from celery.result import AsyncResult
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.decorators import api_view
import json
from django.http import HttpResponse

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches


from backend.celery import app
from .tasks import retina_detection, retina_analyse
from .converting import convert_file_to_image

from .serializers import (
    RetinaFileUploadSerializer,
    RetinaLabelGroupSerializer,
    RetinaLabelMarkSerializer,
    RetinaAnalysisTaskSerializers,
    RetinaAnalysisCorrectResultSerializer
)
from .models import (
    RetinaFileUpload,
    RetinaDetectionTask,
    RetinaLabelGroup,
    RetinaLabelMark,
    RetinaImageCropOCT,
    RetinaAnalysisTask,
    RetinaAnalysisCropTask,
    RetinaAnalysisCorrectResult,
)


class RetinaLabelGroups(APIView):
    def get(self, request, format=None):
        label_groups = RetinaLabelGroup.objects.all()
        serializer = RetinaLabelGroupSerializer(label_groups, many=True)
        return Response(serializer.data)


class RetinaLabelMarks(APIView):
    def get(self, request, format=None):
        label_marks = RetinaLabelMark.objects.all()
        serializer = RetinaLabelMarkSerializer(label_marks, many=True)
        return Response(serializer.data)


class RetinaFileExamples(APIView):
    def get(self, request, format=None):
        examples = RetinaFileUpload.objects.filter(is_example=True)[0:20]
        serializer = RetinaFileUploadSerializer(examples, many=True)
        return Response(serializer.data)


@api_view(["POST"])
def upload_retina_file(request):
    try:
        file = request.data["file"]
        filename = file.name  # достаем оригинальное имя
        extension = filename.split(".")[-1].lower()  # расширение файла
        retina_file = RetinaFileUpload.objects.create(ext=extension)
        retina_file.file_orig.save(file.name, file)

        # convert file to image
        if extension not in ["jpg", "jpeg"]:
            # тут преобразуем если файл не изображение
            image_file = convert_file_to_image(retina_file.file_orig.path)
            retina_file.image.save(image_file.name, image_file)
        else:
            # Копируем оригинал в папку
            retina_file.image.save(file.name, file)

        serializer = RetinaFileUploadSerializer(retina_file, many=False)
        return Response(serializer.data)
    except Exception as e:
        return Response(
            {"error": 500, "message": ("Ошибка сервера: " + str(e))}, status=500
        )


@api_view(["POST"])
def detection_retina_image(request):
    try:
        if request.method == "POST":
            image_guid = request.POST.get("image_guid")
            is_obb = bool(int(request.POST.get("is_obb", 0)))
            image_upload = RetinaFileUpload.objects.get(guid=image_guid)
            retina_detection_task = RetinaDetectionTask(
                image=image_upload, is_obb=is_obb
            )
            retina_detection_task.save()
            task_queue = retina_detection.delay(retina_detection_task.guid)
            return Response(
                {
                    "error": 0,
                    "message": "Задача детекции успешно поставлена в очередь",
                    "data": {
                        "task_model_id": retina_detection_task.guid,
                        "task_queue_id": task_queue.id,
                    },
                },
                status=202,
            )
        else:
            return Response(
                {"error": 400, "message": "Ошибка запроса получения детекции "},
                status=400,
            )
    except Exception as e:
        return Response({"error": 500, "message": ("Ошибка: " + str(e))}, status=500)


@api_view(["GET"])
def get_status_detection_retina_image(request, task_id):
    try:
        task = AsyncResult(task_id, app=app)
        if task.ready():
            response = {
                "task_status": task.state,
                "message": "Изображение успешно распознано",
                "results": {
                    "crops": task.info.get("crops", []),
                    "task": task.info.get("task", {}),
                },
            }
            return Response(response, status=200)
        elif task.state == "SUCCESS":
            # задача выполнена
            response = {
                "task_status": task.state,
                "message": str(task.info.get("status_text", "")),
                "results": {
                    "crops": task.info.get("crops", []),
                    "task": task.info.get("task", {}),
                },
            }
            return Response(response, status=200)
        elif task.state == "PROGRESS":
            # задача в процессе обработки
            response = {
                "task_status": task.state,
                "message": str(task.info.get("status_text", "")),
                "results": {
                    "crops": task.info.get("crops", []),
                    "task": task.info.get("task", {}),
                },
            }
            return Response(response, status=200)
        elif task.state == "PENDING":
            # задача в ожидании
            response = {
                "task_status": task.state,
                "message": "Задача детекции успешно поставлена в очередь",
                "results": {"crops": [], "task": {}},
            }
            return Response(response, status=200)
        elif task.state == "STARTED":
            # ошибка в процессе обработки
            response = {
                "task_status": task.state,
                "message": "Начинаем задачу детекции сегментов изображения",
                "results": {"crops": [], "task": {}},
            }
            return Response(response, status=200)
        elif task.state == "FAILURE":
            # ошибка в процессе обработки
            response = {
                "task_status": task.state,
                "message": str(task.info.get("status_text", "")),
                "results": {"crops": [], "task": {}},
            }
            return Response(response, status=500)
        else:
            # во всех остальных случаях
            response = {
                "task_status": "FAILURE",
                "message": "Задача, поставленная в очередь завершилась с ошибкой",
                "results": {"crops": [], "task": {}},
            }
            return Response(response, status=500)
    except Exception as e:
        return Response(
            {
                "task_status": "FAILURE",
                "message": ("Ошибка: " + str(e)),
                "results": {"crops": [], "task": {}},
            },
            status=500,
        )


@api_view(["POST"])
def analysis_retina_image(request):
    try:
        json_data = json.loads(request.body.decode("utf-8"))
        task_detection_guid = json_data["task_detection"]
        crops = json_data["crops"]
        task_detection = RetinaDetectionTask.objects.get(guid=task_detection_guid)
        crops = RetinaImageCropOCT.objects.filter(
            task_detection=task_detection, pk__in=crops
        )

        if len(crops) == 0:
            raise Exception("Не указаны изображения сетчатки для анализа")

        task_analysis = RetinaAnalysisTask(task_detection=task_detection)
        task_analysis.save()
        for crop in crops:
            task_analysis.crops.add(crop)
        task_analysis.save()
        task_queue = retina_analyse.delay(task_analysis.guid)
        return Response(
            {
                "error": 0,
                "message": "Задача анализа успешно поставлена в очередь",
                "data": {
                    "task_model_id": str(task_analysis.guid),
                    "task_queue_id": task_queue.id,
                },
            },
            status=202,
        )
    except Exception as e:
        return Response({"error": 500, "message": ("Ошибка: " + str(e))}, status=500)


@api_view(["GET"])
def get_status_analysis_retina_image(request, task_id):
    task = AsyncResult(task_id, app=app)
    response = {
        "task_status": task.state,
        "message": str(task.info.get("status_text", "")),
        "progress": task.info.get("progress", False),
        "results": task.info.get("results", []),
    }
    return Response(response, status=200)


@api_view(["POST"])
def correct_task_analysis(request, task_id):
    task_analysis = RetinaAnalysisTask.objects.get(guid=task_id)
    json_data = json.loads(request.body.decode("utf-8"))
    crops_corrected = json_data['data']
    print(crops_corrected)
    # Результаты анализа еще не правили
    for crop in crops_corrected:
        crop_id = crop['crop_id']
        labels = crop['labels']
        crop_analysis = RetinaAnalysisCropTask.objects.get(crop_id=crop_id, task_analysis=task_analysis)
        
        if task_analysis.is_corrected:
            # Если уже исправляем результат, то удаляем все предыдущие значения
            old_results = RetinaAnalysisCorrectResult.objects.filter(retina_analysis_crop_task=crop_analysis)
            old_results.delete()

        for label_id in labels:
            label = RetinaLabelMark.objects.get(id=label_id)
            correctResult = RetinaAnalysisCorrectResult(retina_analysis_crop_task=crop_analysis, label_mark=label)
            correctResult.save()
    
    # todo Добавить проверку что могут передать пустой массив для исправления
    task_analysis.is_corrected = True
    task_analysis.save()

    retina_analyse_serializer = RetinaAnalysisTaskSerializers(task_analysis, many=False)
    retina_analysis_crop_tasks = RetinaAnalysisCropTask.objects.filter(task_analysis=task_analysis)
    corrected_results = RetinaAnalysisCorrectResult.objects.filter(retina_analysis_crop_task__in=retina_analysis_crop_tasks)
    crops_corrected_serializer = RetinaAnalysisCorrectResultSerializer(corrected_results, many=True)

    response = {
        'error': 0,
        'message': 'Изменения внесены в базу данных',
        'task_analysis': retina_analyse_serializer.data,
        'corrected': crops_corrected_serializer.data
    }
    return Response(response, status=200)



@api_view(['POST'])
def generate_report(request, task_analysis_guid):
    try:
        description = request.POST.get('description')
        task_analysis = RetinaAnalysisTask.objects.get(guid=task_analysis_guid)
        task_upload = task_analysis.task_detection.image
        image_path = task_upload.image.path
        doc = Document()
        # Add a Title to the document
        doc.add_heading('Отчет Retina DeepAI',0)
        p = doc.add_paragraph()
        # Выравниваем абзац по середине 
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run('Не является медицинской услугой, служит для интерпретации медицинских изображений при проведении рутинных процессов динамического наблюдения пациентов с патологией макулярной области сетчатки')
        run.font.size = Pt(12)

        # Добавялем картинку
        doc.add_picture(image_path, width=Inches(6))#, height=Inches(2))

        doc.add_paragraph()
        p = doc.add_paragraph()
        # Выравниваем абзац по середине 
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run('Результаты анализа:')
        run.font.size = Pt(16)
        p = doc.add_paragraph()
        run = p.add_run(description)
        run.font.size = Pt(16)

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=download.docx'
        doc.save(response)

        return response
    except Exception as e:
        return Response({'state': 'FAILURE', 'message': ('Ошибка: ' + str(e)), 'progress': False, 'results': []}, status=500)
